"""Stage 1 — Parse: route by file type, return one plaintext blob per résumé.

PDF → PyMuPDF; DOCX → python-docx; text/markdown → decode. Scanned PDFs
(< 100 extracted chars/page) raise ScannedPDFError — the Mistral OCR path is
v0.5 scope, not v0.1.
"""
import logging
from io import BytesIO

from app.config import DOCX_MIME, PDF_MIME, TEXT_MIMES

logger = logging.getLogger(__name__)

SCANNED_CHARS_PER_PAGE = 100


class ParsingError(Exception):
    """Base class — anything that prevents turning the file into text."""


class UnsupportedFileTypeError(ParsingError):
    pass


class ScannedPDFError(ParsingError):
    pass


class EmptyDocumentError(ParsingError):
    pass


def parse_resume(content: bytes, mime_type: str) -> str:
    if mime_type == PDF_MIME:
        text = _parse_pdf(content)
    elif mime_type == DOCX_MIME:
        text = _parse_docx(content)
    elif mime_type in TEXT_MIMES:
        text = content.decode("utf-8", errors="replace")
    else:
        raise UnsupportedFileTypeError(f"unsupported mime type: {mime_type}")

    text = text.strip()
    if not text:
        raise EmptyDocumentError("document contains no extractable text")
    return text


def _parse_pdf(content: bytes) -> str:
    import pymupdf  # lazy: heavy native import

    try:
        doc = pymupdf.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ParsingError(f"could not open PDF: {exc}") from exc

    try:
        pages = [page.get_text() for page in doc]
    finally:
        doc.close()

    text = "\n".join(pages)
    page_count = max(1, len(pages))
    if len(text.strip()) < SCANNED_CHARS_PER_PAGE * page_count:
        raise ScannedPDFError(
            "PDF appears to be scanned (too little extractable text); "
            "OCR is not supported in v0.1"
        )
    return text


def _parse_docx(content: bytes) -> str:
    import docx  # lazy

    try:
        document = docx.Document(BytesIO(content))
    except Exception as exc:
        raise ParsingError(f"could not open DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)
